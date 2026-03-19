"""
Document ingestion pipeline: parsing, chunking, and embedding.
Updated for llama3.2:3b - optimized for speed with larger chunks.

Design rationale:
- Supports PDF, DOCX, and TXT formats (most common for legislative documents)
- RecursiveCharacterTextSplitter preserves document structure better than simple splitting
- Chunk size 2000 chars (≈ 400-500 tokens) - Updated for llama3.2:3b - reduces chunk count from 163 to ~30-40
- Chunk overlap 200 chars - Updated for llama3.2:3b - prevents semantic breaks across boundaries
- Token counting upstream enables compression and efficiency tracking
- Larger chunks work well with llama3.2:3b's efficient context handling
"""

import os
import logging
import tiktoken
from typing import List, Tuple, Dict
from pathlib import Path

# Document parsing
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# LangChain chunking
from langchain_text_splitters import RecursiveCharacterTextSplitter

from vector_store import VectorStore

logger = logging.getLogger(__name__)


class DocumentIngester:
    """Handles parsing, chunking, and embedding of legal documents."""
    
    SUPPORTED_FORMATS = {".pdf", ".docx", ".txt"}
    CHUNK_SIZE = 2000  # Updated for llama3.2:3b - larger chunks = fewer chunks (~30-40 instead of 163)
    CHUNK_OVERLAP = 200  # Updated for llama3.2:3b - increased overlap to preserve context
    
    def __init__(self, vector_store: VectorStore):
        """
        Initialize ingester with vector store for embedding chunks.
        
        Args:
            vector_store: Initialized VectorStore instance
        """
        self.vector_store = vector_store
        
        # Initialize tiktoken for accurate token counting
        # Using "cl100k_base" encoding (GPT-3.5 and GPT-4)
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        
        # RecursiveCharacterTextSplitter tries to keep semantic units together
        # Splits on: "\n\n" (paragraphs), "\n" (lines), " " (words), "" (chars)
        # This preserves document structure better than naive character splitting
        # Updated for llama3.2:3b with larger chunk_size and chunk_overlap
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.CHUNK_SIZE,  # Updated for llama3.2:3b: 2000 chars
            chunk_overlap=self.CHUNK_OVERLAP,  # Updated for llama3.2:3b: 200 chars
            separators=["\n\n", "\n", " ", ""],
            length_function=len,
        )
        
        logger.info("DocumentIngester initialized with optimized chunk settings for llama3.2:3b")
    
    def ingest(self, file_path: str, doc_id: str) -> Dict:
        """
        Main ingestion workflow: parse -> chunk -> embed -> store -> return stats.
        
        Args:
            file_path: Path to document file
            doc_id: Unique document identifier
        
        Returns:
            Dict with ingestion results:
            {
                "doc_id": str,
                "filename": str,
                "chunk_count": int,
                "total_tokens": int,
                "status": "completed" or "failed",
                "error": str (if failed)
            }
        """
        try:
            logger.info(f"Starting ingestion for {file_path}")
            
            # Step 1: Parse document
            text, metadata = self._parse_file(file_path)
            
            # Step 2: Count total tokens in original document
            total_tokens = len(self.tokenizer.encode(text))
            logger.info(f"Document contains {total_tokens} tokens before chunking")
            
            # Step 3: Split into chunks (reduced from 163 to ~30-40 with larger chunk_size for llama3.2:3b)
            chunks = self.splitter.split_text(text)
            logger.info(f"Document split into {len(chunks)} chunks (optimized for llama3.2:3b - expected 30-40 chunks)")
            
            # Step 4: Generate metadata for each chunk
            chunk_metadata = self._generate_chunk_metadata(
                chunks, 
                metadata,
                doc_id
            )
            
            # Step 5: Embed and store chunks
            chunk_count = self.vector_store.add_chunks(
                doc_id=doc_id,
                chunks=chunks,
                metadata=chunk_metadata
            )
            
            result = {
                "doc_id": doc_id,
                "filename": os.path.basename(file_path),
                "chunk_count": chunk_count,
                "total_tokens": total_tokens,
                "status": "completed",
            }
            logger.info(f"Ingestion completed: {result}")
            return result
            
        except Exception as e:
            logger.error(f"Ingestion failed: {e}", exc_info=True)
            return {
                "doc_id": doc_id,
                "filename": os.path.basename(file_path) if file_path else "unknown",
                "status": "failed",
                "error": str(e),
            }
    
    def _parse_file(self, file_path: str) -> Tuple[str, Dict]:
        """
        Parse document based on file extension.
        
        Returns:
            (text_content, metadata_dict)
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {file_ext}. Supported: {self.SUPPORTED_FORMATS}")
        
        if file_ext == ".pdf":
            return self._parse_pdf(file_path)
        elif file_ext == ".docx":
            return self._parse_docx(file_path)
        elif file_ext == ".txt":
            return self._parse_txt(file_path)
    
    def _parse_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from PDF using PyPDF2.
        
        Note: For complex PDFs with images/tables, consider upgrading to PyMuPDF
        """
        try:
            reader = PdfReader(file_path)
            text_parts = []
            page_count = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                # Add minimal formatting to preserve structure
                text_parts.append(f"--- PAGE {page_num} ---\n{page_text}\n")
            
            full_text = "\n".join(text_parts)
            metadata = {
                "source": os.path.basename(file_path),
                "format": "pdf",
                "page_count": page_count,
            }
            return full_text, metadata
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            raise
    
    def _parse_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from DOCX using python-docx."""
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)
            
            metadata = {
                "source": os.path.basename(file_path),
                "format": "docx",
                "paragraph_count": len(doc.paragraphs),
            }
            return full_text, metadata
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise
    
    def _parse_txt(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from plain text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                full_text = f.read()
            
            metadata = {
                "source": os.path.basename(file_path),
                "format": "txt",
            }
            return full_text, metadata
        except Exception as e:
            logger.error(f"TXT parsing failed: {e}")
            raise
    
    def _generate_chunk_metadata(
        self, 
        chunks: List[str], 
        file_metadata: Dict,
        doc_id: str
    ) -> List[Dict]:
        """
        Create metadata for each chunk for traceability and filtering.
        
        Metadata enables:
        - Source document tracking
        - Token counting per chunk (for compression stats)
        - Filtering by section or page
        """
        chunk_metadata = []
        for idx, chunk in enumerate(chunks):
            # Count tokens in this chunk for efficiency tracking
            chunk_tokens = len(self.tokenizer.encode(chunk))
            
            metadata = {
                "doc_id": doc_id,
                "chunk_index": idx,
                "source": file_metadata.get("source", "unknown"),
                "format": file_metadata.get("format", "unknown"),
                "token_count": chunk_tokens,
                "section_title": "general",  # Could be improved with section detection
            }
            chunk_metadata.append(metadata)
        
        return chunk_metadata
    
    def count_tokens(self, text: str) -> int:
        """
        Count tokens in text using tiktoken.
        Exposed as utility for token tracking throughout app.
        """
        return len(self.tokenizer.encode(text))
